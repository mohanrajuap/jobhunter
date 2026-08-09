"""Text extraction from PDF / DOCX / TXT resumes."""

from __future__ import annotations

import logging
from pathlib import Path

log = logging.getLogger(__name__)


def _from_pdf(path: Path) -> str:
    try:
        import pdfplumber
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise RuntimeError("pdfplumber is required to read PDF resumes: pip install pdfplumber") from exc

    chunks: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            chunks.append(page.extract_text() or "")
    return "\n".join(chunks)


def _from_docx(path: Path) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise RuntimeError("python-docx is required to read DOCX resumes: pip install python-docx") from exc

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    # Skills often live in tables, which paragraph iteration misses entirely.
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def extract_text(path: str | Path) -> str:
    """Return the plain text of a resume. Raises on unreadable/unsupported files."""
    resume = Path(path).expanduser()
    if not resume.exists():
        raise FileNotFoundError(f"Resume not found: {resume}")

    suffix = resume.suffix.lower()
    if suffix == ".pdf":
        text = _from_pdf(resume)
    elif suffix in (".docx", ".docm"):
        text = _from_docx(resume)
    elif suffix in (".txt", ".md"):
        text = resume.read_text(encoding="utf-8", errors="ignore")
    elif suffix == ".doc":
        raise RuntimeError(
            f"Legacy .doc is not supported ({resume.name}). Save it as .docx or .pdf and retry."
        )
    else:
        raise RuntimeError(f"Unsupported resume format '{suffix}'. Use PDF, DOCX or TXT.")

    if len(text.strip()) < 100:
        log.warning(
            "Extracted only %d characters from %s — if it is a scanned image, "
            "keyword matching will be poor. Export a text-based PDF instead.",
            len(text.strip()), resume.name,
        )
    return text
